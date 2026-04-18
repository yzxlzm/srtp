from operator import itemgetter
from typing import List, Optional

import os
import tempfile
from pathlib import Path

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnableLambda
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import TextLoader

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from src.llm import get_llm
from src.embeddings import get_embeddings
from src.vectorstore import build_vectorstore, load_vectorstore, save_vectorstore


def load_docx(file_path: Path) -> List[Document]:
    """使用 python-docx 加载 Word 文档"""
    if DocxDocument is None:
        raise ImportError("需要安装 python-docx: pip install python-docx")
    
    doc = DocxDocument(str(file_path))
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    
    content = "\n".join(paragraphs)
    if not content.strip():
        return []
    
    return [Document(page_content=content, metadata={"source": str(file_path)})]


def load_docs(path: Optional[str] = None):
    """
    加载文档，支持 .txt, .docx, .md 格式
    如果 path 为 None，则加载 data 目录下的所有支持格式的文档
    """
    base_dir = Path(__file__).resolve().parent.parent
    data_dir = base_dir / "data"
    
    all_docs = []
    
    if path:
        # 加载指定文件
        file_path = Path(path)
        if not file_path.is_absolute():
            file_path = base_dir / file_path
        
        if file_path.suffix.lower() == ".docx":
            docs = load_docx(file_path)
            all_docs.extend(docs)
        else:
            loader = TextLoader(str(file_path), encoding="utf-8")
            all_docs.extend(loader.load())
    else:
        # 加载 data 目录下的所有支持格式的文档
        supported_extensions = {".txt", ".md", ".docx"}
        
        for file_path in sorted(data_dir.iterdir()):  # 排序以便按顺序加载
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                try:
                    if file_path.suffix.lower() == ".docx":
                        docs = load_docx(file_path)
                    else:
                        loader = TextLoader(str(file_path), encoding="utf-8")
                        docs = loader.load()
                    all_docs.extend(docs)
                    print(f"✓ 已加载文档: {file_path.name} (共 {len(docs)} 个文档块)")
                except Exception as e:
                    print(f"✗ 加载文档 {file_path.name} 时出错: {e}")
    
    if not all_docs:
        raise ValueError("未找到任何可加载的文档")
    
    print(f"\n开始处理文档，共 {len(all_docs)} 个原始文档块...")
    # 加大分块，减少文本块数量以加快构建向量库
    splitter = RecursiveCharacterTextSplitter(chunk_size=900, chunk_overlap=80)
    split_docs = splitter.split_documents(all_docs)
    print(f"文档分割完成，共生成 {len(split_docs)} 个文本块\n")
    return split_docs


def docs_to_context(docs):
    """将检索到的文档转换为上下文字符串"""
    if not docs:
        return ""
    return "\n\n".join(d.page_content for d in docs)


def check_relevance(docs_with_scores, threshold=0.6):
    """
    检查检索到的文档是否与问题相关
    FAISS 返回的是距离分数（L2距离），越小越相似
    对于 bge-small 模型，通常相关文档的距离 < 0.6
    返回: (is_relevant, context)
    """
    if not docs_with_scores:
        return False, ""
    
    # 检查最小距离（最相似的文档）
    min_distance = min(score for _, score in docs_with_scores)
    
    # 如果最小距离大于阈值，认为不相关（距离越大越不相似）
    # 降低阈值到 0.6，使判断更严格
    if min_distance > threshold:
        return False, ""
    
    # 提取相关文档内容（距离小于阈值的文档，最多3个）
    relevant_docs = [doc for doc, score in docs_with_scores if score <= threshold][:3]
    context = "\n\n".join(d.page_content for d in relevant_docs)
    return True, context


def build_rag_chain():
    base_dir = Path(__file__).resolve().parent.parent

    # 缓存目录：可通过环境变量 VECTOR_CACHE_DIR 指定；默认放在系统临时目录下，避免中文/空格路径引起的 FAISS 读写问题
    env_cache = os.environ.get("VECTOR_CACHE_DIR")
    if env_cache:
        vector_dir = Path(env_cache)
    else:
        vector_dir = Path(tempfile.gettempdir()) / "vectorstore_cache"

    vector_dir.mkdir(parents=True, exist_ok=True)

    embeddings = get_embeddings()

    # 优先加载已有向量库，避免每次重建
    index_path = vector_dir  # 使用 ASCII 路径，避免 faiss 在中文/空格目录下报错
    index_f = index_path / "index.faiss"
    index_pkl = index_path / "index.pkl"

    try:
        if index_f.exists() and index_pkl.exists():
            vectorstore = load_vectorstore(index_path, embeddings)
            print(f"已加载缓存向量库：{index_path}")
        else:
            raise FileNotFoundError("缓存文件缺失")
    except Exception as e:
        print(f"缓存向量库不可用，开始重新构建: {e}")
        docs = load_docs()
        vectorstore = build_vectorstore(docs, embeddings)
        try:
            save_vectorstore(vectorstore, index_path)
            print(f"向量库已构建并缓存到：{index_path}")
        except Exception as se:
            print(f"向量库缓存失败（继续使用内存向量库）：{se}")

    # 配置检索器：k=5 以获取更多候选
    retriever = vectorstore.as_retriever(search_kwargs={"k": 5})
    llm = get_llm()

    # RAG 提示词：当有相关上下文时使用
    rag_prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                """你是一个专业的高校人才供需适配系统助手。

重要规则：
1. 如果上下文信息能够回答用户的问题，请直接基于上下文给出准确、专业的回答。
2. 如果上下文信息无法回答用户的问题（例如：问题涉及未来时间、上下文没有相关信息等），请直接使用你的知识回答，绝对不要说"根据上下文信息，无法获取"、"上下文未涉及"等拒绝性话语。
3. 绝对不要在回答中提及"上下文"、"根据上下文"、"给定的上下文"、"提供的上下文信息"等措辞。
4. 不要补充来源说明或引用说明。
5. 回答要专业、友好、简洁，直接给出答案。

上下文：{context}""",
            ),
            ("human", "问题：{question}"),
        ]
    )

    # 直接 LLM 提示词：当没有相关上下文时使用
    direct_prompt = ChatPromptTemplate.from_messages(
        [
            (
                "system",
                "你是一个专业的高校人才供需适配系统助手。请直接回答用户的问题，回答要专业、友好、简洁。",
            ),
            ("human", "问题：{question}"),
        ]
    )

    def route_question(inputs):
        """根据检索结果决定使用 RAG 还是直接 LLM"""
        question = inputs["question"]
        
        # 使用 similarity_search_with_score 获取带分数的检索结果
        docs_with_scores = vectorstore.similarity_search_with_score(question, k=5)
        
        # 检查相关性（阈值 0.6，FAISS 的 L2 距离越小越相似）
        # 对于 bge-small 模型，相关文档通常距离 < 0.6
        # 降低阈值使判断更严格，避免不相关问题进入 RAG 模式
        is_relevant, context = check_relevance(docs_with_scores, threshold=0.6)
        
        if is_relevant and context.strip():
            # 有相关上下文，使用 RAG
            return rag_prompt.invoke({"question": question, "context": context})
        else:
            # 无相关上下文，直接使用 LLM
            return direct_prompt.invoke({"question": question})
    
    chain = (
        {"question": itemgetter("question")}
        | RunnableLambda(route_question)
        | llm
    )
    return chain

